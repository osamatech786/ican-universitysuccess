import streamlit as st
from datetime import datetime, date
from streamlit_drawable_canvas import st_canvas
import json
import pandas as pd
from docx import Document
from docx.shared import Inches
import io
from PIL import Image
import numpy as np
import smtplib
from email.message import EmailMessage
import re
from dotenv import load_dotenv
import os

# Set page configuration with a favicon
st.set_page_config(
    page_title="ICAN Brunel University",
    page_icon="https://www.brunel.ac.uk/_MB34Eg_746fb155-c770-4062-ae0e-ece2f206213c/static-main/img/brunel-logo.png",  # Path to your logo
    layout="centered"  # "centered" or "wide"
)

# add render support along with st.secret
def get_secret(key):
    try:
        load_dotenv()
        # Attempt to get the secret from environment variables
        secret = os.environ.get(key)
        if secret is None:
            raise ValueError("Secret not found in environment variables")
        return secret
    except (ValueError, TypeError) as e:
        # If an error occurs, fall back to Streamlit secrets
        if hasattr(st, 'secrets'):
            return st.secrets.get(key)
        # If still not found, return None or handle as needed
        return None

if 'files' not in st.session_state:
    st.session_state.files = []

# Load country names and dialing codes from the JSON file
with open("resources/world-countries.json") as file:
    data = json.load(file)
    countries = {entry['name']: entry['dialing_code'] for entry in data}  # Map country name to dialing code
country_names = ["Select"] + sorted(countries.keys())  # Sort country names

# Load subject areas from the text file
with open("resources/subject_area_list.txt", "r") as file:
    subject_areas = [line.strip() for line in file.readlines()]

# The sub-options for each subject area, like Foundation, Undergraduate, etc.
sub_options = [
    "Foundation",
    "Undergraduate",
    "Pre-Masters",
    "Postgraduate",
    "PhD & Research",
    "Professional development"
]

# Learning modes (these are hypothetical, but you can adjust them to your case)
learning_modes = ["Online", "Blended", "On-Campus"]

# Function to validate the phone number
def validate_phone_number(phone, dialing_code):
    # Remove all spaces and dashes from the phone number
    phone = phone.replace(" ", "").replace("-", "")
    
    # Check if the phone number starts with the correct dialing code
    if not phone.startswith(dialing_code):
        return False, f"Phone number must start with {dialing_code}."
    
    # Extract the number part (remove the dialing code)
    number_without_code = phone[len(dialing_code):]
    
    # Ensure the number part contains only digits and has a valid length (e.g., 10-15 digits)
    if not number_without_code.isdigit():
        return False, "Phone number must contain only digits after the dialing code."
    
    if not (10 <= len(number_without_code) <= 15):
        return False, "Phone number must be between 10 and 15 digits long (excluding country code)."
    
    return True, ""

def is_valid_email(email):
    # Comprehensive regex for email validation
    pattern = r'''
        ^                         # Start of string
        (?!.*[._%+-]{2})          # No consecutive special characters
        [a-zA-Z0-9._%+-]{1,64}    # Local part: allowed characters and length limit
        (?<![._%+-])              # No special characters at the end of local part
        @                         # "@" symbol
        [a-zA-Z0-9.-]+            # Domain part: allowed characters
        (?<![.-])                 # No special characters at the end of domain
        \.[a-zA-Z]{2,}$           # Top-level domain with minimum 2 characters
    '''
    
    # Match the entire email against the pattern
    return re.match(pattern, email, re.VERBOSE) is not None

def is_signature_drawn(signature):
    # Check if signature is None or an empty numpy array
    if signature is None:
        return False
    # Ensure it is a numpy array and has content
    if isinstance(signature, np.ndarray) and signature.size > 0:
        # Additional check: if the array is not just empty white pixels
        # Assuming white background is [255, 255, 255] in RGB
        if np.all(signature == 255):
            return False
        return True
    return False

# Function to send email with attachments (Handle Local + Uploaded)
def send_email_with_attachments(sender_email, sender_password, receiver_email, subject, body, files=None, local_file_path=None):
    msg = EmailMessage()
    msg['From'] = sender_email
    msg['To'] = ", ".join(receiver_email)
    msg['Subject'] = subject
    msg.set_content(body, subtype='html')

    # Attach uploaded files
    if files:
        for uploaded_file in files:
            uploaded_file.seek(0)  # Move to the beginning of the UploadedFile
            msg.add_attachment(uploaded_file.read(), maintype='application', subtype='octet-stream', filename=uploaded_file.name)

    # Attach local file if specified
    if local_file_path:
        with open(local_file_path, 'rb') as f:
            file_data = f.read()
            file_name = local_file_path.split('/')[-1]
            msg.add_attachment(file_data, maintype='application', subtype='octet-stream', filename=file_name)

    # Use Gmail SMTP server for sending the email (office365 for outlook)
    with smtplib.SMTP('smtp.gmail.com', 587) as server:
        server.ehlo()
        server.starttls()
        server.login(sender_email, sender_password)
        server.send_message(msg)

# Initialize session state variables if they do not exist
if 'step' not in st.session_state:
    st.session_state.step = 1
    st.session_state.submission_done = False
    st.session_state.personal_info = ""  # Full name
    st.session_state.dob = None  # Date of birth
    st.session_state.gender = "Select"  # Gender
    st.session_state.country = ""  # Country of residence
    st.session_state.email = ""  # Email address
    st.session_state.phone = ""  # Phone number
    st.session_state.address = ""  # Residential address
    st.session_state.previous_qualifications = ""  # Previous qualifications
    st.session_state.current_institution = ""  # Current institution
    # st.session_state.start_date = None  # Uncomment if needed
    st.session_state.front_id_document = None  # Front ID document
    st.session_state.back_id_document = None  # Back ID document
    st.session_state.address_proof = None  # Address proof document
    st.session_state.additional_document = None  # Additional documents if needed
    st.session_state.learning_preferences = ""  # Learning preferences
    st.session_state.special_requirements = ""  # Special requirements
    st.session_state.emergency_contact = ""  # Emergency contact information
    st.session_state.consent = False  # Consent for data processing
    st.session_state.signature = None  # Store signature


# Define a function to calculate progress and percentage
def get_progress(step, total_steps=14):
    return int((step / total_steps) * 100)


# Define the total number of steps
total_steps = 14

# Calculate the current progress
progress = get_progress(st.session_state.step, total_steps)

# Display the progress bar and percentage
st.write(f"Progress: {progress}%")
st.progress(progress)


# Define the different steps
if st.session_state.step == 1:
    st.image('resources/AspireCraft_resized.gif', use_column_width=True)
    # st.image(Image.open('resources/logo.png').resize((500, 300)), use_column_width=True)

    st.title("WELCOME TO ASPIRECRAFT!")
    st.write("""
    At AspireCraft, we believe in unlocking potential and creating opportunities for lifelong learning.
    Our international CPD and accredited qualifications are designed to empower you with the skills and knowledge needed to excel in your chosen field.

    We are excited to have you on board and look forward to supporting your journey towards achieving UK accreditation.

    Let's get started with your enrolment process. It's simple and straightforward. Please proceed by filling out the following fields one at a time.
    Click 'Next' to begin your journey with AspireCraft!
    """)
    if st.button("Next"):
        st.session_state.step = 2
        st.experimental_rerun()

elif st.session_state.step == 2:
    st.title("> 1: Personal Information")
    
    # Ensure the personal_info variable is correctly set from the session state
    st.session_state.personal_info = st.text_input(
        "Please enter your full name as it appears on your official documents.",
        value=st.session_state.personal_info  # Retain previous value
    )

    # Next and Back buttons for navigation
    next_clicked = st.button("Next", key=f"next_{st.session_state.step}")
    back_clicked = st.button("Back", key=f"back_{st.session_state.step}")

    # Handle Next button click
    if next_clicked:
        if st.session_state.personal_info:  # Check if the field is not empty
            st.session_state.step = 3  # Move to the next step
            st.experimental_rerun()  # Refresh the app to reflect the new step
        else:
            st.warning("Please enter your full name before proceeding.")

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 1  # Go back to the previous step (Section 1)
        st.experimental_rerun()  # Refresh to update the step


elif st.session_state.step == 3:
    st.title("> 2: Date of Birth")
    # Check if dob is a string and convert it back to a date object
    if isinstance(st.session_state.get("dob"), str):
        st.session_state.dob = datetime.strptime(st.session_state.get("dob"), "%d-%m-%Y").date()

    # Date of Birth
    st.session_state.dob = st.date_input(
        label="Date of Birth",  # Label for the field
        value=st.session_state.get("dob"),  # Correctly access dob from session state
        min_value=date(1900, 1, 1),  # Minimum selectable date
        max_value=date.today(),  # Maximum selectable date
        help="Choose a date",  # Tooltip text
        format='DD/MM/YYYY'
    )
    
    # Next and Back buttons for navigation
    next_clicked = st.button("Next", key=f"next_{st.session_state.step}")
    back_clicked = st.button("Back", key=f"back_{st.session_state.step}")

    # Handle Next button click
    if next_clicked:
        if st.session_state.dob:
            # Convert the selected date to the desired string format (DD-MM-YYYY) only when proceeding to the next step
            # st.session_state.dob = st.session_state.dob.strftime("%d-%m-%Y")

            st.session_state.step = 4
            st.experimental_rerun()
        else:
            st.warning("Please select your date of birth before proceeding.")

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 2  # Go back to the previous step (Section 1)
        st.experimental_rerun()

elif st.session_state.step == 4:
    st.title("> 3: Gender")

    # Initialize gender if it doesn't exist
    if 'gender' not in st.session_state:
        st.session_state.gender = "Select"  # Default value

    # Select gender using the selectbox, retaining the previous value
    st.session_state.gender = st.selectbox(
        "Please select your gender.", 
        ["Select", "Male", "Female", "Other"],
        index=["Select", "Male", "Female", "Other"].index(st.session_state.gender)  # Set default value based on session state
    )

    # Next and Back buttons for navigation
    next_clicked = st.button("Next", key=f"next_{st.session_state.step}")
    back_clicked = st.button("Back", key=f"back_{st.session_state.step}")

    # Handle Next button click
    if next_clicked:
        if st.session_state.gender != "Select":
            st.session_state.step = 5
            st.experimental_rerun()
        else:
            st.warning("Please select your gender before proceeding.")

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 3  # Go back to the previous step (Section 2)
        st.experimental_rerun()


# Step 5: Country selection
if st.session_state.step == 5:
    st.title("> 4: Country")

    # Initialize country if it doesn't exist
    if 'country' not in st.session_state:
        st.session_state.country = "Select"  # Default value

    # Select country using the selectbox, retaining the previous value
    st.session_state.country = st.selectbox(
        "Please select your country.", 
        country_names,
        index=country_names.index(st.session_state.country) if st.session_state.country in country_names else 0  # Set default value based on session state
    )

    # Next and Back buttons for navigation
    next_clicked = st.button("Next", key=f"next_{st.session_state.step}")
    back_clicked = st.button("Back", key=f"back_{st.session_state.step}")

    # Handle Next button click
    if next_clicked:
        if st.session_state.country != "Select":
            st.session_state.step = 6
            st.experimental_rerun()
        else:
            st.warning("Please select your country before proceeding.")

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 4  # Go back to the previous step (Section 3)
        st.experimental_rerun()

# Step 6: Contact Information
elif st.session_state.step == 6:
    st.title("> 5: Contact Information")

    # Get the selected country's dialing code from the countries dictionary
    selected_dialing_code = countries.get(st.session_state.country, "")

    # Initialize fields if they do not exist
    if 'email' not in st.session_state:
        st.session_state.email = ""  # Default to empty string
    if 'phone' not in st.session_state:
        st.session_state.phone = ""  # Default to empty string
    if 'address' not in st.session_state:
        st.session_state.address = ""  # Default to empty string

    # Input fields for contact information
    st.session_state.email = st.text_input("Please enter your email address where we can reach you.", value=st.session_state.email)

    # Display the country dialing code before the phone number input
    st.session_state.phone = st.text_input(
        f"Please enter your WhatsApp number (international format starting with {selected_dialing_code} for {st.session_state.country}):", 
        value=st.session_state.phone
    )

    # Display the WhatsApp call availability message
    st.markdown(
        """
        ### Ensure WhatsApp Call Availability:
        We may contact you via WhatsApp. Please make sure your phone number is connected to WhatsApp and can receive international calls.
        """
    )
    # Display clickable images in a single line
    st.write("Download WhatsApp for your device:")

    st.markdown(
        """
        <div style="display: flex; justify-content: space-around; align-items: center;">
            <a href="https://play.google.com/store/apps/details?id=com.whatsapp" target="_blank">
                <img src="https://raw.githubusercontent.com/osamatech786/ican-universitysuccess/refs/heads/main/resources/icons/android.png" alt="Download for Android" style="width:100px;height:100px;margin:10px;">
            </a>
            <a href="https://apps.apple.com/app/whatsapp-messenger/id310633997" target="_blank">
                <img src="https://cdn3.iconfinder.com/data/icons/social-media-logos-i-filled-line/2048/5315_-_Apple-512.png" alt="Download for iOS" style="width:100px;height:100px;margin:10px;">
            </a>
            <a href="https://get.microsoft.com/installer/download/9NKSQGP7F2NH" target="_blank">
                <img src="https://github.com/osamatech786/ican-universitysuccess/blob/main/resources/icons/windows.png?raw=true" alt="Download for Windows" style="width:100px;height:100px;margin:10px;">
            </a>
            <a href="https://web.whatsapp.com/desktop/mac_native/release/?configuration=Release" target="_blank">
                <img src="https://github.com/osamatech786/ican-universitysuccess/blob/main/resources/icons/macbook.png?raw=true" alt="Download for Mac" style="width:100px;height:100px;margin:10px;">
            </a>
        </div>
        """, 
        unsafe_allow_html=True
    )

    # Input for address
    st.session_state.address = st.text_area(
        "Please enter your complete mailing address.", 
        value=st.session_state.address
    )

    # Next and Back buttons for navigation
    next_clicked = st.button("Next", key=f"next_{st.session_state.step}")
    back_clicked = st.button("Back", key=f"back_{st.session_state.step}")

    # Handle Next button click
    if next_clicked:
        if st.session_state.phone and st.session_state.email and st.session_state.address:
            if is_valid_email(st.session_state.email):
                is_valid, message = validate_phone_number(st.session_state.phone, selected_dialing_code)
                if is_valid:
                    st.session_state.step = 7
                    st.experimental_rerun()
                else:
                    st.warning(message)
            else:
                st.warning("Please enter a valid email address.")
        else:
            st.warning("Please fill out all the contact information fields before proceeding.")

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 5  # Go back to the previous step (Country selection)
        st.experimental_rerun()

elif st.session_state.step == 7:
    st.title("> 6: Educational Background")

    # Initialize fields if they do not exist
    if 'previous_qualifications' not in st.session_state:
        st.session_state.previous_qualifications = ""  # Default to empty string
    if 'current_institution' not in st.session_state:
        st.session_state.current_institution = ""  # Default to empty string

    # Input fields with default values from session state
    st.session_state.previous_qualifications = st.text_area(
        "Please list your previous qualifications.", 
        value=st.session_state.previous_qualifications
    )
    st.session_state.current_institution = st.text_input(
        "Please enter the name of your current educational institution (if applicable, else put 'none').", 
        value=st.session_state.current_institution
    )

    # Navigation buttons
    next_clicked = st.button("Next", key=f"next_{st.session_state.step}")
    back_clicked = st.button("Back", key=f"back_{st.session_state.step}")

    # Handle Next button click
    if next_clicked:
        if (st.session_state.previous_qualifications.strip() and 
            st.session_state.current_institution.strip()):
            st.session_state.step = 8
            st.experimental_rerun()
        else:
            st.warning("Please list your previous qualifications and current institution before proceeding.")

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 6  # Go back to the previous step (Section 5)
        st.experimental_rerun()


# Step 7: Course Information (with Subject Areas, Sub-options, and Learning Modes)
if st.session_state.step == 8:
    st.title("> 7: Course Information")

    # Initialize subject area, sub-option, and learning mode if they do not exist
    if 'subject_area' not in st.session_state:
        st.session_state.subject_area = "Select"  # Default value
    if 'sub_option' not in st.session_state:
        st.session_state.sub_option = "Select"  # Default value
    if 'learning_mode' not in st.session_state:
        st.session_state.learning_mode = "Select"  # Default value
    if 'selected_course' not in st.session_state:
        st.session_state.selected_course = {}  # To store selected subject area, course level, and learning mode

    # Subject area selection
    st.session_state.subject_area = st.selectbox(
        "Please select the subject area.", 
        ["Select"] + sorted(subject_areas),  # Subject areas loaded from the text file
        index=(sorted(subject_areas).index(st.session_state.subject_area) + 1) if st.session_state.subject_area in subject_areas else 0
    )

    # Sub-option selection based on the selected subject area
    st.session_state.sub_option = st.selectbox(
        "Please select your course level.", 
        ["Select"] + sub_options,
        index=(sub_options.index(st.session_state.sub_option) + 1) if st.session_state.sub_option in sub_options else 0
    )

    # Learning mode selection
    st.session_state.learning_mode = st.selectbox(
        "Please select the learning mode.", 
        ["Select"] + learning_modes,
        index=(learning_modes.index(st.session_state.learning_mode) + 1) if st.session_state.learning_mode in learning_modes else 0
    )


    # Navigation buttons
    next_clicked = st.button("Next", key=f"next_{st.session_state.step}")
    back_clicked = st.button("Back", key=f"back_{st.session_state.step}")

    # Handle Next button click
    if next_clicked:
        if (st.session_state.subject_area != "Select" 
            and st.session_state.sub_option != "Select"
            and st.session_state.learning_mode != "Select"):
            
            # Store the selected subject area, course level, and learning mode
            st.session_state.selected_course = {
                'subject_area': st.session_state.subject_area,
                'course_level': st.session_state.sub_option,
                'learning_mode': st.session_state.learning_mode
            }
            
            st.session_state.step = 11  # Move to the next step
            st.experimental_rerun()
        else:
            st.warning("Please select the subject area, course level, and learning mode before proceeding.")

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 7  # Go back to the previous step (Section 6)
        st.experimental_rerun()

elif st.session_state.step == 9:
    st.title("> 8: Identification Documents")
    st.text("(*Upload of any 1 document is mandatory)")

    # Upload front and back of the document
    st.session_state.front_id_document = st.file_uploader("Please upload a scan or photo of the front of your passport or ID.", type=["jpg", "png", "pdf", "docx"], key="front")
    if st.session_state.front_id_document is not None:
        if st.session_state.front_id_document not in st.session_state.files:
            st.session_state.files.append(st.session_state.front_id_document)
    # if st.session_state.front_id_document is not None:
    #     st.session_state.files(st.session_state.front_id_document)

    st.session_state.back_id_document = st.file_uploader("Please upload a scan or photo of the back of your passport or ID.", type=["jpg", "png", "pdf", "docx"], key="back")
    if st.session_state.back_id_document is not None:
        if st.session_state.back_id_document not in st.session_state.files:
            st.session_state.files.append(st.session_state.back_id_document)
    # if st.session_state.back_id_document is not None:
    #      st.session_state.files(st.session_state.back_id_document)
    
    # Navigation buttons
    next_clicked = st.button("Next", key=f"next_{st.session_state.step}")
    back_clicked = st.button("Back", key=f"back_{st.session_state.step}")

    # Handle Next button click
    if next_clicked:
        if st.session_state.front_id_document or st.session_state.back_id_document:
            st.session_state.step = 10
            st.experimental_rerun()
        else:
            st.warning("Please upload both the front and back of your identification document before proceeding.")

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 8  # Go back to the previous step (Section 7)
        st.experimental_rerun()

elif st.session_state.step == 10:
    st.title("> 9: Proof of Address")
    st.session_state.address_proof = st.file_uploader("*Please upload a scan or photo of your proof of address.", type=["jpg", "png", "pdf", "docx"])
    if st.session_state.address_proof is not None:
        if st.session_state.address_proof not in st.session_state.files:
            st.session_state.files.append(st.session_state.address_proof)

    # Navigation buttons
    next_clicked = st.button("Next", key=f"next_{st.session_state.step}")
    back_clicked = st.button("Back", key=f"back_{st.session_state.step}")

    # Handle Next button click
    if next_clicked:
        if st.session_state.address_proof:
            st.session_state.step = 11
            st.experimental_rerun()
        else:
            st.warning("Please upload your proof of address before proceeding.")

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 9  # Go back to the previous step (Section 8)
        st.experimental_rerun()

elif st.session_state.step == 11:
    st.title("> 10: Additional Information")

    # Initialize fields if they do not exist
    if 'learning_preferences' not in st.session_state:
        st.session_state.learning_preferences = ""  # Default to empty string
    if 'special_requirements' not in st.session_state:
        st.session_state.special_requirements = ""  # Default to empty string
    if 'emergency_contact' not in st.session_state:
        st.session_state.emergency_contact = ""  # Default to empty string
    if 'consent' not in st.session_state:
        st.session_state.consent = False  # Default to unchecked

    # Input fields with default values from session state
    st.session_state.learning_preferences = st.text_area(
        "Please describe any learning preferences you have.", 
        value=st.session_state.learning_preferences
    )
    st.session_state.special_requirements = st.text_area(
        "Please let us know if you have any special requirements.", 
        value=st.session_state.special_requirements
    )
    st.session_state.emergency_contact = st.text_input(
        "Please provide emergency contact details.", 
        value=st.session_state.emergency_contact
    )
    st.session_state.consent = st.checkbox(
        "I consent to the collection and processing of my personal data according to AspireCraft’s privacy policy.", 
        value=st.session_state.consent
    )

    # Link to the privacy policy
    privacy_policy_doc_link = 'https://drive.google.com/file/d/1QnmwPyUv22LPOU3eKBT1ho55QW_5_olS/view'
    st.write(f"[Privacy Policy]({privacy_policy_doc_link})")  # Actual link to privacy policy

    # Path to the PDF file in the resources folder
    # pdf_file_path = os.path.join('resources', 'Student Privacy Notice_30.07.2024_Rev.1_FF.pdf')
    # Display the link for the PDF file to open in a new tab
    # st.markdown(f'<a href="file://{pdf_file_path}" target="_blank">Privacy Policy</a>', unsafe_allow_html=True)

    # Navigation buttons
    next_clicked = st.button("Next", key=f"next_{st.session_state.step}")
    back_clicked = st.button("Back", key=f"back_{st.session_state.step}")

    # Handle Next button click
    if next_clicked:
        if all([st.session_state.learning_preferences, st.session_state.special_requirements, st.session_state.emergency_contact, st.session_state.consent]):
            st.session_state.step = 12
            st.experimental_rerun()
        else:
            st.warning("Please complete all fields and consent before proceeding.")

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 8  # Go back to the previous step (Section 9)
        st.experimental_rerun()


elif st.session_state.step == 12:
    st.title("> 11: Signature")
    st.write("Please provide your signature below:")

    canvas_result = st_canvas(
        stroke_width=2,
        stroke_color="black",
        background_color="white",
        update_streamlit=True,
        height=150,
        width=600,
        drawing_mode="freedraw",
        key="signature_canvas"
    )
    # Only update the session state if there is a change in the canvas
    if canvas_result.image_data is not None:
        st.session_state.signature = canvas_result.image_data

    # Navigation buttons
    next_clicked = st.button("Next", key=f"next_{st.session_state.step}")
    back_clicked = st.button("Back", key=f"back_{st.session_state.step}")

    # Handle Next button click
    if next_clicked:
        if is_signature_drawn(st.session_state.signature):
        # if st.session_state.signature is not None:
            st.session_state.step = 13
            st.experimental_rerun()
        else:
            st.warning("Please provide your signature before proceeding.")

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 11  # Go back to the previous step (Section 10)
        st.experimental_rerun()

elif st.session_state.step == 13:
    st.title("Final Review")
    st.write("Thank you for providing your details. Please review your information and click 'Submit' to complete your enrolment.")

    st.write(f"**Full Name:** {st.session_state.personal_info}")
    dob = st.session_state.dob.strftime('%d-%m-%Y')
    st.write(f"**Date of Birth:** {dob}")
    st.write(f"**Gender:** {st.session_state.gender}")
    st.write(f"**Country:** {st.session_state.country}")
    st.write(f"**Email:** {st.session_state.email}")
    st.write(f"**Phone:** {st.session_state.phone}")
    st.write(f"**Address:** {st.session_state.address}")
    st.write(f"**Previous Qualifications:** {st.session_state.previous_qualifications}")
    st.write(f"**Current Institution:** {st.session_state.current_institution}")

    # Update this to show the selected course
    if 'selected_course' in st.session_state and st.session_state.selected_course:
        course_info = f"{st.session_state.selected_course['subject_area']} - {st.session_state.selected_course['course_level']} ({st.session_state.selected_course['learning_mode']})"
    else:
        course_info = "None"
    st.write(f"**Course Interested In:** {course_info}")

    # st.write(f"**Learning Mode:** {st.session_state.learning_mode}")
    st.write(f"**Learning Preferences:** {st.session_state.learning_preferences}")
    st.write(f"**Special Requirements:** {st.session_state.special_requirements}")
    st.write(f"**Emergency Contact:** {st.session_state.emergency_contact}")

    if st.session_state.signature is not None:
        st.image(st.session_state.signature, caption="Your Signature")

    # Print the list of files
    if st.session_state.files:
        st.write("Files uploaded:", len(st.session_state.files))
        for file in st.session_state.files:
            st.write(f"File name: {file.name}, File type: {file.type}")
    else:
        st.write("No files uploaded.")
    
    # Submit button
    submit_clicked = st.button("Submit")


###############################
    with st.spinner('Processing....'):

        # Handle Submit button click
        if submit_clicked:        
            # Create a new Document
            doc = Document()
            doc.add_heading('Enrolment Form Submission', 0)

            # Add form details
            doc.add_paragraph(f'Full Name: {st.session_state.personal_info}')

            dob = st.session_state.dob.strftime('%d-%m-%Y')
            doc.add_paragraph(f'Date of Birth: {dob}')

            doc.add_paragraph(f'Gender: {st.session_state.gender}')
            doc.add_paragraph(f'Country: {st.session_state.country}')
            doc.add_paragraph(f'Email: {st.session_state.email}')
            doc.add_paragraph(f'Phone: {st.session_state.phone}')
            doc.add_paragraph(f'Address: {st.session_state.address}')
            doc.add_paragraph(f'Previous Qualifications: {st.session_state.previous_qualifications}')
            doc.add_paragraph(f'Current Institution: {st.session_state.current_institution}')
            
            # Add selected course information
            if 'selected_course' in st.session_state and st.session_state.selected_course:
                course_info = f"{st.session_state.selected_course['subject_area']} - {st.session_state.selected_course['course_level']} ({st.session_state.selected_course['learning_mode']})"
                doc.add_paragraph(f'Course Interested In: {course_info}')
            else:
                doc.add_paragraph(f'Course Interested In: None')

            # doc.add_paragraph(f'Preferred Start Date: {st.session_state.start_date.strftime("%d-%m-%Y")}')

            # doc.add_paragraph(f'Learning Mode: {st.session_state.learning_mode}')
            doc.add_paragraph(f'Learning Preferences: {st.session_state.learning_preferences}')
            doc.add_paragraph(f'Special Requirements: {st.session_state.special_requirements}')
            doc.add_paragraph(f'Emergency Contact: {st.session_state.emergency_contact}')

            # Save the signature image if available
            if st.session_state.signature is not None:
                # Convert numpy array to PIL image
                image_data = st.session_state.signature
                image = Image.fromarray(image_data.astype(np.uint8))  # Ensure correct data type
                
                # Save the image to an in-memory file
                image_stream = io.BytesIO()
                image.save(image_stream, format='PNG')
                image_stream.seek(0)
                
                # Add image to docx
                doc.add_picture(image_stream, width=Inches(2))
            
            # Save the document
            doc_path = f"ICAN_Form_Submission_{st.session_state.personal_info}.docx"
            doc.save(doc_path)

            # Email
            # Sender email credentials
            # Credentials: Streamlit host st.secrets
            # sender_email = st.secrets["sender_email"]
            # sender_password = st.secrets["sender_password"]

            sender_email = get_secret("sender_email")
            sender_password = get_secret("sender_password")

            # Credentials: Local env
            # load_dotenv()                                     # uncomment import of this library!
            # sender_email = os.getenv('EMAIL')
            # sender_password = os.getenv('PASSWORD')
            team_email = ["enquiry.aspirecraft@gmail.com"]

            learner_email = [st.session_state.email]
            
            subject_team = f"ICAN - Course: {course_info} Country: {st.session_state.country} Name: {st.session_state.personal_info} Submission Date: {date.today()}"
            body_team = "ICAN Form submitted. Please find attached file."

            subject_learner = "Thank You for Your Interest in Our Courses!"
            body_learner = f"""
            <html>
            <body>
                <p>Dear {st.session_state.personal_info},</p>

                <p>Thank you for expressing your interest in Brunel University courses. AspireCraft is delighted to assist you through our International Career Advice and Navigation (ICAN) service. A member of our team will be contacting you within the next 24 hours to guide you through the next steps of the enrolment process and support your career education.</p>

                <p><strong>What’s Next?</strong></p>

                <ol>
                    <li><strong>Initial Screening and Eligibility Assessment:</strong></li>
                </ol>
                <ul>
                    <li>One of our representatives will call you soon via WhatsApp to conduct a quick screening test to assess your eligibility and the entry requirements for the course. Please ensure that the number you have provided is linked to your phone so we can reach you without any issues.</li>
                    <li>During the call, we will ask a few questions to determine if you meet the criteria to move forward with the admissions process.</li>
                    <li>If you are successful, you will be invited to attend an online session with the University admissions staff. In this session, you can ask any questions about the course, the university, or your career prospects. Afterward, you will be provided with a link to begin the formal enrollment process.</li>
                </ul>

                <p>This call is an important step to ensure that you have the right foundation to succeed in your studies and to provide you with the information you need to feel confident moving forward.</p>

                <p>If you have any immediate questions, feel free to contact us at <strong><a href="mailto:enquiry.aspirecraft@gmail.com">enquiry.aspirecraft@gmail.com</a></strong>.</p>

                <p>We look forward to speaking with you soon and welcoming you to our learning community!</p>

                <p>Best regards,</p>
                <p>Student Admissions Team<br>
                AspireCraft<br>
                <em>CRAFTING SUCCESS, EMPOWERING FUTURES</em></p>
            </body>
            </html>
            """


            # Send email to team with attachments
            if st.session_state.files or doc_path:
                send_email_with_attachments(sender_email, sender_password, team_email, subject_team, body_team, st.session_state.files, doc_path)
            
            # Send thank you email to learner
            send_email_with_attachments(sender_email, sender_password, learner_email, subject_learner, body_learner)

            # Update session state to show the final thank you message
            st.session_state.submission_done = True
            st.session_state.step = 14  # Move to the final step to show the thank you message
            st.experimental_rerun()

#111111111111111111
    # Add a warning before the back button
    st.info("If you go back, you will have to re-sign the form.")

    # Navigation buttons
    back_clicked = st.button("Back", disabled=st.session_state.submission_done)

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 12  # Go back to the previous step
        st.experimental_rerun()
#11111111111111111

# Add a new step for the thank you message
elif st.session_state.step == 14:
    st.title("Thank You!")
    st.write("Check your email for the final boarding.")
    st.write('')
    st.image('resources/logo_org.png', use_column_width=True)

# else:
#     st.write("Form completed. Thank you!")

# streamlit run app.py --server.port 8503
# Dev : https://linkedin.com/in/osamatech786
